from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
import sys
import random
from docx2pdf import convert
import os
from pdf2image import convert_from_path
import PyPDF2



# function prepares text for further processing 
def parse_text_file(text_file):
   
    # read the text_file
    with open(text_file, 'r', encoding = "utf-8") as file:
        text_content = file.read()
    
    # split text into paragraphs
    paragraphs = text_content.split(b'\t'.decode("utf-8"))

    rep_pars = []
    # remove ParlaMint words - are harmful to training process
    for paragraph in paragraphs:
        par = ""
        words = paragraph.split()
        for i in range(len(words)):
            if(not("ParlaMint" in words[i])):
                par += words[i]
                if(i < len(words) - 1):
                    par += " "
        rep_pars.append(par)

    return rep_pars


# randomly assigns bold, italic, underline or colour to the given percentage of words
def random_assign(run, i, boldy, italicy, underliny, colorful):
    if(i in boldy): 
        run.bold = True
    if(i in italicy):
        run.italic = True
    if(i in underliny):
        run.font.underline = True
    if(i in colorful):
        a = random.randrange(0, 2)
       # only allow black, gray and brown colors (reality!)
        if(a == 0):
            color = random.randrange(60)
            run.font.color.rgb = RGBColor(color, color, color)
        else:
            colors = [(51, 43, 39), (46, 19, 14), (38, 20, 0), (23, 9, 3), (43, 29, 17)]
            color = random.choice(colors)
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
    return True
    
    # parameters: [paragraph num, font, size, alignment, columns, bold, italic, underlined, colored]
    #                    0          1     2        3        4       5      6          7        8
def fill_section(doc, parameters, offset, table):
    # if offset == 0, then we are at the first section
    if(offset > 0):
        doc.add_section(WD_SECTION.CONTINUOUS)
    #check the columns number (0 -> 1, 1 -> 2)
    if(parameters[4] >= 1):
        doc.sections[-1]._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'2')       
    else:
        doc.sections[-1]._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'1')
    # set margins
    doc.sections[-1].left_margin = Cm(1.5)
    doc.sections[-1].right_margin = Cm(1.5)
    doc.sections[-1].top_margin = Cm(2.5)
    doc.sections[-1].bottom_margin = Cm(2.5)   
    for i in range(offset, offset + parameters[0]):            
        if(i >= len(table)):
            break
        # new paragraph
        p = doc.add_paragraph()
        p.paragraph_format.alignment = parameters[3]                             
        lines = table[i].split("\n")
        if((i == offset) & (offset > 0)):
            p.add_run("\n")

        for j in range(len(lines)):
            words = lines[j].split()
            l = len(words)
            #prepare the distribution        
            num_bold = l * parameters[5] // 100
            num_italic = l * parameters[6] // 100
            num_underlined = l * parameters[7] // 100
            num_colourful = l * parameters[8] // 100
                    
            bolded = random.sample(range(l), num_bold)
            italiced = random.sample(range(l), num_italic)
            underlined = random.sample(range(l), num_underlined)
            colored = random.sample(range(l), num_colourful)
                    
            for x in range(l):
                run = None
                run = p.add_run(words[x])
                font = run.font
                font.name = parameters[1]
                font.size = Pt(parameters[2])
                random_assign(run, x, bolded, italiced, underlined, colored)  
                if(x < l - 1):
                    p.add_run(" ")
            if(j < len(lines) - 1):
                p.add_run("\n")
    return True

def extract_text(doc, pdf):
    text_by_page = []
    pages = []
    for i in pdf.pages:
        pages.append(i.extract_text())
    
    paragraphs = []
    for i in doc.paragraphs:
        paragraphs.append(i.text)

    start = 0
    cur_par = 0
    broken = 0
    indx = 0
    # going through words (not chars)
    for page in pages:
        text = ""
        if(cur_par >= len(paragraphs)):
            break
        broken_page = "".join("".join(page.split()).split())
        par = paragraphs[cur_par].split()
        indx = 0
        end = False
        if(broken > 0):
            for i in range(broken, len(paragraphs[cur_par].split())):
                ind = broken_page.find(par[i], indx)
                if(ind >= indx):
                    text += par[i]
                    text += " "
                    indx = ind
                else:
                    broken = i
                    text_by_page.append(text)
                    end = True
                    break
            if(not(end)): 
                text += "\n"
                cur_par += 1
                broken = 0
        if(end):
            continue
        if(cur_par >= len(paragraphs)):
            text_by_page.append(text)
            break
        # before examining others: start = cur_par
        start = cur_par
        par = "".join(paragraphs[cur_par].split())
        while(par in broken_page):
            # add paragraph
            indx = broken_page.find(par) + len(par)
            cur_par += 1
            if(cur_par >= len(paragraphs)):
                break
            par = "".join(paragraphs[cur_par].split())
        # add all that are certainly
        for i in range(start, cur_par):
            text += paragraphs[i]
            text += "\n\n"
        text_by_page.append(text)
        if(cur_par >= len(paragraphs)):
            # gone trough all paragraphs
            break
        # check if only part of paragraph is on page
        text = ""
        par = paragraphs[cur_par].split()
        for i in range(len(par)):
            ind = broken_page.find(par[i], indx)
            if(ind >= indx):
                text += par[i]
                text += " "
                indx = ind
            else:
                broken = i
                text_by_page[-1] += text
                break
        start = cur_par
    return text_by_page    

################################################################################################################################
############################################################  MAIN  ############################################################
################################################################################################################################

# Check if command-line arguments are all provided
if len(sys.argv) < 2:
    print("Please provide a text file as an argument.")
    sys.exit(1)
if len(sys.argv) < 3:
    print("Please provide a name for output file as an argument.")
    sys.exit(1)
if len(sys.argv) < 4:
    print("Please provide a number of sections wanted in output.")
    sys.exit(1)
if len(sys.argv) < 13 | (len(sys.argv) - 4 % 9 != 0):
    print("Please provide a correct number of inputs for each section (paragraphs, font, size, alignment, columns 1/2, bold, italic, undelined, colored).")
# Get the text file path from the command-line argument
text_file = sys.argv[1]

# Parse the text file
table = parse_text_file(text_file)
outputName = sys.argv[2]
sectionsNum = int(sys.argv[3])
sectionInfo = []
## contents: list of lists [paragraph num, font, size, alignment, columns, bold, italic, underlined, colored]
for i in range(sectionsNum):
    sectionInfo.append([int(sys.argv[4 + i * 9]), sys.argv[5 + i * 9], int(sys.argv[6 + i * 9]), sys.argv[7 + i * 9], int(sys.argv[8 + i * 9]), int(sys.argv[9 + i * 9]), int(sys.argv[10 + i * 9]), int(sys.argv[11 + i * 9]), int(sys.argv[12 + i * 9])])

##create offset table for number of paragraphs   --- on i-th place is the index of first paragraph that is in i-th section

## possible errors:
# there is more paragraphs listed than given: program slices the given list of parameters maximum of desired sections but will cut the first that goes over len(table) - 1
# there is less paragraphs listed: program connects all the last paragraphs to last section 
offset = []
for i in range(sectionsNum):
    if (i == 0):
        offset.append(0)
    else:
        offset.append(offset[i - 1] + sectionInfo[i - 1][0])
        if offset[i] > (len(table) - 1):
            ## new section should not exist
            sectionsNum = i ## i goes from 0 to secNum - 1
            del offset[i]
            sectionInfo[i-1][0] = len(table) - offset[i - 1]
            print("Section " + str(i + 1) + " is unvalid, it has been corrected in the program; only " + str(i) + " sections now.")
            break
if(offset[-1] + sectionInfo[-1][0] < len(table) - 1):
    sectionInfo[-1][0] = len(table) - offset[-1]

# possible layout parameters
# fonts: 8 
fonts = ["Calibri", "Cambria", "Arial", "Times New Roman", "Comic Sans MS", "Helvetica", "Cambria", "Verdana"]

# alignments: 5
alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.DISTRIBUTE, WD_ALIGN_PARAGRAPH.JUSTIFY]

two_columns = False
for i in range(sectionsNum):
    
    # check size 
    if(sectionInfo[i][2] <= 0 | sectionInfo[i][2] > 70):
        sectionInfo[i][2] = 12
    
    # change alignment string to proper object
    if( sectionInfo[i][3] == "left"):
        sectionInfo[i][3] = alignments[0]
    elif(sectionInfo[i][3] == "center"):
        sectionInfo[i][3] = alignments[1]
    elif(sectionInfo[i][3] == "right"):
        sectionInfo[i][3] = alignments[2]
    elif(sectionInfo[i][3] == "distribute"):
        sectionInfo[i][3] = alignments[3]
    elif(sectionInfo[i][3] == "justify"):
        sectionInfo[i][3] = alignments[4]
    else:
        print("Provided invalid alignment for section " + str(i + 1) + "; applying left alignment.")
        sectionInfo[i][3] = alignments[0]
    
    # check bold, italic, underlined, color
    for j in range(5, 9):
        if(sectionInfo[i][j] < 0 | sectionInfo[i][j] > 100):
            sectionInfo[i][j] = 50
    
    #print("parameters checked")
    if((sectionInfo[i][4] == 1) & two_columns):
        sectionInfo[i][4] = 0
        two_columns = False
    elif(sectionInfo[i][4] == 1):
        two_columns = True
    else:
        two_columns = False
    
# columns: 0 = one column, 1 = two columns
# if the sections is on it ignores columns sections

doc = Document()


# here we want to fill the document based on input data, section by section
for i in range(sectionsNum):
    # Check if the parameters are within posibble inputs
    parameters = sectionInfo[i]
    # Apply to the section and add section to the document
    fill_section(doc, parameters, offset[i], table)

path = os.path.dirname(text_file)
path = os.path.join(path, sys.argv[2])
# creating directory for output files
if(not os.path.exists(path)):
    os.mkdir(path)
pathName = os.path.join(path, sys.argv[2] + ".docx")
pathNamePDF = os.path.join(path, sys.argv[2] + ".pdf")
doc.save(pathName)

# convert word document to PDF
convert(pathName, pathNamePDF)

pdf = open(pathNamePDF, "rb")
pdf = PyPDF2.PdfReader(pdf)
# allign text by pages (word only knows about sections and paragraphs, not about text by page)
pages_text = extract_text(doc, pdf)

# write text by page into txt files
for i in range(len(pages_text)):
    name = os.path.join(path, str(i + 1) + "_page.txt")
    with open(name, "w", encoding="utf-8") as file:
        file.write(pages_text[i])
